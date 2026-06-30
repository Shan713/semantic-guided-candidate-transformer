from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document

from main import main


def test_cli_runs_end_to_end_with_csv_and_resume():
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        resume_path = temp_dir_path / "resume.docx"
        output_path = temp_dir_path / "output.json"
        config_path = temp_dir_path / "projection.yml"

        config_path.write_text(
            """
fields:
  - path: full_name
    type: string
  - path: primary_email
    from: emails[0]
  - path: phone
    from: phones[0]
    normalize: E164
  - path: skills
    from: skills[].name
    normalize: canonical
include_confidence: true
include_provenance: true
on_missing: null
""",
            encoding="utf-8",
        )

        document = Document()
        document.add_paragraph("John Doe")
        document.add_paragraph("john.doe@example.com")
        document.add_paragraph("+1 415 555 2671")
        document.add_paragraph("LinkedIn: https://www.linkedin.com/in/john-doe")
        document.add_paragraph("Skills: Python, Docker")
        document.save(resume_path)

        exit_code = main(
            [
                "--csv",
                str(Path("sample_inputs") / "candidate.csv"),
                "--resume",
                str(resume_path),
                "--config",
                str(config_path),
                "--output",
                str(output_path),
            ]
        )

        assert exit_code == 0
        assert output_path.exists()
        payload = json.loads(output_path.read_text(encoding="utf-8"))
        assert payload["full_name"]
        assert payload["primary_email"]
        assert payload["phone"]
